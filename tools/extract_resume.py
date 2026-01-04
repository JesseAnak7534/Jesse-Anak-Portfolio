import sys
import re
import json
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Missing dependency: python-docx. Install with: python -m pip install python-docx", file=sys.stderr)
    sys.exit(2)

HEADING_ALIASES = {
    'summary': {"summary", "profile", "about", "objective"},
    'experience': {"experience", "work experience", "professional experience", "employment"},
    'education': {"education", "academics"},
    'skills': {"skills", "technical skills", "key skills", "core skills"},
    'projects': {"projects", "personal projects", "selected projects", "portfolio"},
    'certifications': {"certifications", "certificates", "licenses"},
    'awards': {"awards", "honors", "achievements"},
    'languages': {"languages"},
    'interests': {"interests", "hobbies"},
    'publications': {"publications"},
}

# Flatten alias map for detection
ALIAS_TO_CANON = {alias: canon for canon, aliases in HEADING_ALIASES.items() for alias in aliases}

EMAIL_RE = re.compile(r"[\w\.-]+@[\w\.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"(\+?\d[\d\s\-\(\)]{7,}\d)")
URL_RE = re.compile(r"https?://[^\s)]+", re.I)

BULLET_PREFIXES = ("•", "- ", "– ", "— ", "* ")


def read_docx_text(path: Path):
    doc = Document(str(path))
    lines = []
    # Paragraphs
    for p in doc.paragraphs:
        text = (p.text or '').strip()
        if text:
            lines.append(text)
    # Tables (some resumes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or '').strip()
                if text:
                    for subline in [s.strip() for s in text.splitlines() if s.strip()]:
                        lines.append(subline)
    # De-dup consecutive duplicates
    cleaned = []
    last = None
    for line in lines:
        if line != last:
            cleaned.append(line)
        last = line
    return cleaned


def detect_headings(lines):
    # Return list of tuples (index, canonical_heading)
    results = []
    for i, raw in enumerate(lines):
        t = raw.strip().strip(':').lower()
        t_simple = re.sub(r"[^a-z ]", "", t)
        if t in ALIAS_TO_CANON:
            results.append((i, ALIAS_TO_CANON[t]))
        elif t_simple in ALIAS_TO_CANON:
            results.append((i, ALIAS_TO_CANON[t_simple]))
        elif t.isupper() and len(t) <= 30 and t in ALIAS_TO_CANON:
            results.append((i, ALIAS_TO_CANON[t]))
    # Deduplicate keeping first of same index
    seen = set()
    uniq = []
    for idx, name in results:
        if idx not in seen:
            uniq.append((idx, name))
            seen.add(idx)
    # Sort by index
    uniq.sort(key=lambda x: x[0])
    return uniq


def extract_contact(lines):
    contact = {}
    all_text = "\n".join(lines)
    email = EMAIL_RE.search(all_text)
    if email:
        contact['email'] = email.group(0)
    phone = PHONE_RE.search(all_text)
    if phone:
        contact['phone'] = phone.group(1).strip()
    urls = URL_RE.findall(all_text)
    for u in urls:
        lu = u.lower()
        if 'linkedin.com' in lu:
            contact['linkedin'] = u
        elif 'github.com' in lu:
            contact['github'] = u
        elif 'x.com' in lu or 'twitter.com' in lu:
            contact['twitter'] = u
        elif 'behance.net' in lu or 'dribbble.com' in lu:
            contact['portfolio'] = u
        else:
            contact.setdefault('website', u)
    # Location heuristic
    for l in lines[:8]:
        if any(k in l.lower() for k in ['based in', 'location', 'address', 'lagos', 'abuja', 'nigeria', 'state', 'city']):
            contact['location'] = l
            break
    return contact


def split_sections(lines, headings):
    sections = {}
    if not headings:
        return sections
    for idx, (h_i, name) in enumerate(headings):
        start = h_i + 1
        end = headings[idx + 1][0] if idx + 1 < len(headings) else len(lines)
        body = [l for l in lines[start:end] if l.strip()]
        sections[name] = body
    return sections


def extract_name_title(lines):
    # First 3-5 lines often hold name/title
    head = [l for l in lines[:6] if l.strip()]
    name = None
    title = None
    for i, l in enumerate(head):
        if '@' in l or 'http' in l.lower() or 'linkedin.com' in l.lower():
            continue
        # Likely name if it has 2-4 words, capitalized
        words = l.split()
        caps_ratio = sum(w[:1].isupper() for w in words) / max(1, len(words))
        if 1 < len(words) <= 5 and caps_ratio >= 0.6:
            name = l
            # next line as title if short
            if i + 1 < len(head) and len(head[i+1]) <= 60:
                title = head[i+1]
            break
    return name, title


def parse_bullets(lines):
    items = []
    for l in lines:
        s = l.strip()
        for b in BULLET_PREFIXES:
            if s.startswith(b):
                s = s[len(b):].strip()
                break
        if s:
            items.append(s)
    return items


def parse_skills(lines):
    text = " ".join(lines)
    # Split by commas or bullets or pipes
    raw = re.split(r"[\n\r,•|]+", text)
    skills = []
    for s in raw:
        t = s.strip()
        if not t:
            continue
        # Filter out long phrases
        if len(t.split()) > 6:
            continue
        skills.append(t)
    # Deduplicate preserving order
    seen = set()
    uniq = []
    for s in skills:
        k = s.lower()
        if k not in seen:
            uniq.append(s)
            seen.add(k)
    return uniq


def chunk_entries(lines):
    # Split by blank lines or strong bullet markers
    entries = []
    buf = []
    for l in lines:
        if not l.strip():
            if buf:
                entries.append(buf)
                buf = []
            continue
        if l.strip().startswith(BULLET_PREFIXES) and buf:
            # treat as continuation bullet
            buf.append(l)
        else:
            buf.append(l)
    if buf:
        entries.append(buf)
    return entries


def parse_experience(lines):
    entries = chunk_entries(lines)
    result = []
    for e in entries:
        head = e[0] if e else ''
        role = None
        company = None
        dates = None
        # Common patterns: "Role - Company - Dates" or "Role at Company (Dates)"
        if ' at ' in head:
            parts = head.split(' at ', 1)
            role = parts[0].strip(' -–—')
            rest = parts[1]
            m = re.search(r"\(([^)]+)\)$", rest)
            if m:
                company = rest[:m.start()].strip(' -–—')
                dates = m.group(1).strip()
            else:
                company = rest.strip()
        else:
            parts = re.split(r"\s[–—-]\s", head)
            if len(parts) >= 2:
                role, company = parts[0].strip(), parts[1].strip()
            if len(parts) >= 3:
                dates = parts[2].strip()
        bullets = parse_bullets(e[1:]) if len(e) > 1 else []
        result.append({
            'role': role or head,
            'company': company,
            'dates': dates,
            'details': bullets if bullets else e[1:],
        })
    return result


def parse_education(lines):
    entries = chunk_entries(lines)
    result = []
    for e in entries:
        head = e[0]
        degree = None
        school = None
        dates = None
        parts = re.split(r"\s[–—-]\s", head)
        if len(parts) >= 2:
            degree, school = parts[0].strip(), parts[1].strip()
        if len(parts) >= 3:
            dates = parts[2].strip()
        result.append({
            'degree': degree or head,
            'school': school,
            'dates': dates,
            'details': e[1:],
        })
    return result


def parse_projects(lines):
    entries = chunk_entries(lines)
    result = []
    for e in entries:
        title = e[0]
        bullets = parse_bullets(e[1:]) if len(e) > 1 else []
        url = None
        m = URL_RE.search(" ".join(e))
        if m:
            url = m.group(0)
        result.append({
            'title': title,
            'url': url,
            'details': bullets if bullets else e[1:],
        })
    return result


def main():
    if len(sys.argv) < 3:
        print("Usage: extract_resume.py <input.docx> <output.json>")
        sys.exit(1)
    in_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    lines = read_docx_text(in_path)
    headings = detect_headings(lines)
    sections = split_sections(lines, headings)

    name, title = extract_name_title(lines)
    contact = extract_contact(lines)

    data = {
        'name': name,
        'title': title,
        'contact': contact,
        'summary': "\n".join(sections.get('summary', [])[:5]) if sections.get('summary') else None,
        'skills': parse_skills(sections.get('skills', [])) if sections.get('skills') else None,
        'experience': parse_experience(sections.get('experience', [])) if sections.get('experience') else None,
        'education': parse_education(sections.get('education', [])) if sections.get('education') else None,
        'projects': parse_projects(sections.get('projects', [])) if sections.get('projects') else None,
        'certifications': sections.get('certifications'),
        'awards': sections.get('awards'),
        'languages': sections.get('languages'),
        'interests': sections.get('interests'),
        'publications': sections.get('publications'),
        'fullText': lines,
        'sectionsRaw': sections,
    }

    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open('w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"Wrote JSON to {out_path}")


if __name__ == '__main__':
    main()
